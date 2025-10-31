# Updated Streamlit app to reflect new content structure and clean errors
import streamlit as st
import pandas as pd
import openai
from dotenv import load_dotenv
import os
import json

load_dotenv()
openai.api_key = os.getenv("OPENAI_API_KEY")

st.set_page_config(page_title="Product Listing Automation", layout="wide")
st.title("📦 Product Listing Automation Tool")

# Function to generate SEO keywords automatically
def generate_seo_keywords(product_name, key_benefits, description_summary, brand_tonality):
    """
    Calls the OpenAI API to generate a structured JSON array of SEO keywords.
    """
    if not openai.api_key:
        st.error("❌ OpenAI API Key is not set.")
        return []

    # Define the System Instruction
    system_instruction = (
        "You are an expert SEO keyword researcher specializing in e-commerce (Amazon, Shopify, etc.). "
        "Your sole task is to analyze the provided product name, key benefits, and description to "
        "generate a list of 25 high-value, relevant SEO keywords. Keywords must be relevant to "
        "purchase intent, long-tail, and address common customer search queries. "
        "Output MUST be a single JSON object containing an array. "
        "JSON Structure: {\"keywords\": [{\"keyword\": \"exact keyword phrase\", \"type\": \"Primary, Long-Tail, Competitor, or LSI\", \"search_intent\": \"Commercial, Transactional, or Informational\"}, ...]}"
    )

    # Define the User Prompt
    user_prompt = f"""
Generate the SEO keywords JSON for the following product:

1. Product Name: {product_name}
2. Core USPs/Key Benefits: {key_benefits}
3. Product Description Summary: {description_summary}
4. Brand Context (Tonality): {brand_tonality}
"""
    
    try:
        response = openai.chat.completions.create(
            model="gpt-4o-mini", 
            messages=[
                {"role": "system", "content": system_instruction},
                {"role": "user", "content": user_prompt}
            ],
            response_format={"type": "json_object"},
            temperature=0.7
        )
        
        json_output = response.choices[0].message.content
        parsed_json = json.loads(json_output)
        
        # Extract the keywords list
        if isinstance(parsed_json, list):
            keywords_list = parsed_json
        elif isinstance(parsed_json, dict):
            keywords_list = parsed_json.get('keywords', []) or parsed_json.get('data', [])
        else:
            st.error("❌ Failed to parse structured JSON output from API.")
            return []

        return keywords_list

    except json.JSONDecodeError:
        st.error("❌ API returned invalid JSON. Please retry.")
    except Exception as e:
        st.error(f"❌ OpenAI API Error: {str(e)}")
        
    return []

# Brand Tonality Dictionary - Embedded in code
BRAND_TONALITIES = {
    "MakeMeeBold": """Confident, not cocky: Customers are sure about our products because they work, not because of overselling.
Modern & aspirational: Sleek, smart, and relevant to today's beauty consumer — informed but not clinical.
Empowering: Encourages self-expression, not perfection. The name itself speaks to inner and outer confidence.
Clean & refined: Language is crisp and minimal, avoiding clutter or fluff.
Contemporary feminine: It's a women-focused brand, however it is gender-neutral enough to feel inclusive, but still elegant and aspirational.
Results-oriented: clear about efficacy, with evidence, or focus on USPs""",

    "Urban Yog": """Playful, not pretentious: Skincare is fun and expressive; never boring or overly serious
Relatable, not preachy: Speaks like a best-friend who gets it - conversational, and easy to connect with
Bold & unapologetic: owns topics that others avoid such as body hair, acne, etc with confidence, humour and zero shame.
Yummy & sensory: language is colourful & indulgent - textures, adjectives, scents, and visuals that make skincare feel like a treat 
Youthful & vibrant: fresh, pop-culture-driven, socially fluent, resonates with Gen Z and young millenials
Witty & confident: Uses playful quips & puns sometimes, that add spunk & personality without losing meaning or clarity.
Miniso, cutesy vibes: The brand gives off a miniso-esqe, cutesy vibe""",

    "Urban Gabru": """Confident, not arrogant: Speaks to men who take charge of their lives — ambitious, self-assured, and results-driven.
Sharp & modern: Sleek, innovative, and tuned to the needs of the contemporary Indian man.
Empowering & aspirational: Encourages self-improvement and personal growth — grooming is positioned as the start of greatness, not the end goal.
Practical & purposeful: Clear, direct messaging that communicates product efficacy and everyday relevance — built to support hustle, fitness, and professional ambitions.
Masculine, yet approachable: Maintains a strong, confident tone without being intimidating — relatable to men from all walks of life.
Dynamic & motivating: Language inspires action and upward momentum.
Premium & credible: the brand positions itself as a premium and credible brand with authenticity and trustworthiness.
To-the-point: Clearly defined USPs and no extra fluff just the way most men communicate.""",

    "Seoulskin": """Calm, not flashy: Focuses on simplicity and serenity, avoiding hype and quick-fix messaging.
Consistent & reliable: Emphasizes long-term results and daily rituals, positioning the brand as a trustworthy skincare companion.
Purposeful & minimal: Every product and message is intentional — no clutter, no unnecessary complexity, only what truly benefits the skin.
Gentle & nurturing: Soft, caring language conveys comfort and reassurance, appealing to those seeking mindful self-care.
Authentic & credible: Rooted in the principles of Korean skincare, the brand communicates expertise without being clinical or intimidating.
Sophisticated & understated: Elegant, refined, and approachable — appeals to consumers who value quality and efficacy over trends.
Mindful & reassuring: Speaks to the desire for slow, consistent care — emphasizing trust, comfort, and skin confidence."""
}

# Helper function to normalize and match brand names
def normalize_brand_name(brand_name):
    """
    Normalize brand name for case-insensitive and spacing-insensitive matching.
    Returns the standardized brand name from BRAND_TONALITIES keys.
    """
    if not brand_name:
        return None
    
    # Clean the input: strip, lowercase, remove extra spaces
    cleaned = ' '.join(brand_name.strip().lower().split())
    
    # Create a mapping of normalized names to actual keys
    brand_mapping = {
        'makemeebold': 'MakeMeeBold',
        'make mee bold': 'MakeMeeBold',
        'urbanyog': 'Urban Yog',
        'urban yog': 'Urban Yog',
        'urbangabru': 'Urban Gabru',
        'urban gabru': 'Urban Gabru',
        'seoulskin': 'Seoulskin',
        'seoul skin': 'Seoulskin'
    }
    
    # Check direct mapping first
    if cleaned in brand_mapping:
        return brand_mapping[cleaned]
    
    # Try to match by removing all spaces
    cleaned_no_space = cleaned.replace(' ', '')
    for key, value in brand_mapping.items():
        if cleaned_no_space == key.replace(' ', ''):
            return value
    
    # Try partial matching with BRAND_TONALITIES keys
    for brand_key in BRAND_TONALITIES.keys():
        if cleaned == brand_key.lower() or cleaned.replace(' ', '') == brand_key.lower().replace(' ', ''):
            return brand_key
    
    return None

st.subheader("📥 Upload KLD Sheet")
uploaded_file = st.file_uploader("Upload your KLD Excel file", type=["xlsx"])

if uploaded_file:
    try:
        # Read Excel file
        df = pd.read_excel(uploaded_file, header=None)
        
        # Extract columns 3 and 4 (index 3, 4) which are Particulars and Details in new format
        df_clean = df.iloc[1:, [3, 4]]
        df_clean.columns = ['Particulars', 'Details']
        
        # Remove rows where both Particulars and Details are empty
        df_clean = df_clean.dropna(subset=['Particulars', 'Details'], how='all')
        
        # Remove rows where Particulars is empty but keep rows where only Details might be empty
        df_clean = df_clean[df_clean['Particulars'].notna()]
        
        # Strip whitespace from column values
        df_clean['Particulars'] = df_clean['Particulars'].astype(str).str.strip()
        df_clean['Details'] = df_clean['Details'].astype(str).str.strip()
        
        # Remove rows where Particulars is empty string after stripping
        df_clean = df_clean[df_clean['Particulars'] != '']
        
        # Create a dictionary for easier field matching
        product_data = {}
        for _, row in df_clean.iterrows():
            key = row['Particulars']
            value = row['Details'] if pd.notna(row['Details']) and row['Details'] != 'nan' else ''
            product_data[key] = value
        
        # Helper function to get field with flexible matching (case-insensitive and partial match)
        def get_field(*possible_keys, default="Not specified"):
            """Try multiple possible keys and return first match"""
            for key in possible_keys:
                # Exact match (case-insensitive)
                for data_key, value in product_data.items():
                    if key.lower() == data_key.lower() and value:
                        return value
                # Partial match (case-insensitive)
                for data_key, value in product_data.items():
                    if key.lower() in data_key.lower() and value:
                        return value
            return default
        
        # Create normalized product_info dictionary with flexible field mapping
        product_info = pd.Series({
            'Product Name': get_field('Product Name'),
            'Brand Name': get_field('Brand Name'),
            'USPs Front': get_field('USPs Front', 'USP Front'),
            'USP Back': get_field('USP Back'),
            'USP Side': get_field('USP Side', default=''),
            'Ingredients': get_field('INGREDIENTS', 'Ingredients'),
            'Claims': get_field('Claims'),
            'How to use it?': get_field('HOW TO USE IT?', 'HOW TO USE', 'How to use it?', 'How to use'),
            'MRP (Incl. of all taxes)': get_field('MRP (Incl. of all taxes)', 'MRP'),
            'Category': get_field('Category', 'Product Type', 'Application Area', default=''),
            'Target Audience': get_field('Target Audience', 'Ideal For', default=''),
            'Ideal For (Gender)': get_field('Ideal For (Gender)', 'Ideal for (Gender)', 'Ideal For', 'Gender', default=''),
            'KNOW YOUR PRODUCT': get_field('KNOW YOUR PRODUCT', 'Know Your Product', default=''),
            'Net Qty.': get_field('Net Qty.', 'Net Qty', default=''),
            'Country Of Origin': get_field('Country Of Origin', 'Country of Origin', default=''),
            'Brand Owned & Marketed By': get_field('BRAND OWNED & MARKETED BY', 'Brand Owned & Marketed By', 'Marketed By', default=''),
            'Email': get_field('EMAIL', 'Email', 'E-mail', default=''),
            'Contact Us': get_field('CONTACT US', 'Contact Us', 'Contact', default=''),
            'Box Includes': get_field('BOX INCLUDES', 'Box Includes', 'Box includes', 'What\'s in the box', default=''),
            'Warranty': get_field('WARRANTY', 'Warranty', 'Warrenty', default=''),
        })
        
        st.success("✅ KLD sheet loaded and parsed successfully!")
        
    except Exception as e:
        st.error(f"❌ Error parsing file: {str(e)}")
        st.info("💡 Please ensure your Excel file has 'Particulars' in column D and 'Details' in column E")
        st.stop()

    # Category Selector - Mandatory field
    st.subheader("📂 Select Product Category")
    
    # Try to get category from the file first
    detected_category = product_info.get('Category', '')
    
    # Determine if category was actually detected from sheet (not just the default 'Beauty')
    # Check if it's a meaningful detection by looking at the original field
    category_detected = detected_category and detected_category != 'Beauty'
    
    # Category selection with radio buttons
    if category_detected:
        # If category detected from sheet, auto-select it
        if 'Electronics' in str(detected_category):
            default_index = 1
        elif 'Beauty' in str(detected_category):
            default_index = 0
        else:
            default_index = None
        
        selected_category = st.radio(
            "Choose the product category (Required):",
            options=["Beauty", "Electronics"],
            index=default_index,
            horizontal=True,
            help="Category detected from sheet. You can change if needed."
        )
    else:
        # If no category detected, don't pre-select anything - let user choose
        selected_category = st.radio(
            "Choose the product category (Required):",
            options=["Beauty", "Electronics"],
            index=None,
            horizontal=True,
            help="Please select the product category to continue"
        )
    
    # Show warning if no category is selected
    if not selected_category:
        st.warning("⚠️ Please select a product category to continue")
        st.stop()
    
    # Update the product_info with selected category
    product_info['Category'] = selected_category
    
    st.success(f"✅ Category set to: **{selected_category}**")
    
    # Brand Selector - Mandatory field
    st.subheader("🏷️ Select Brand")
    
    # Try to detect brand from the file
    detected_brand_raw = product_info.get('Brand Name', '')
    detected_brand = normalize_brand_name(detected_brand_raw)
    
    # Get available brands from tonality file
    available_brands = list(BRAND_TONALITIES.keys()) if BRAND_TONALITIES else []
    
    if available_brands:
        # Set default index based on normalized brand detection
        default_index = 0
        if detected_brand and detected_brand in available_brands:
            default_index = available_brands.index(detected_brand)
        
        selected_brand = st.selectbox(
            "Choose the brand (Required):",
            options=available_brands,
            index=default_index,
            help="This determines the tone and style of generated content"
        )
        
        # Show the tonality for selected brand
        if selected_brand and selected_brand in BRAND_TONALITIES:
            with st.expander("📖 View Brand Tonality Guidelines", expanded=False):
                st.markdown(f"**{selected_brand} Tonality:**")
                st.text(BRAND_TONALITIES[selected_brand])
        
        st.success(f"✅ Brand set to: **{selected_brand}**")
    else:
        st.warning("⚠️ No brand tonality file found. Content will be generated with generic tone.")
        selected_brand = None
    
    st.markdown("---")

    st.subheader("🔹 Product Information")
    
    # Display extracted product information in a clean table
    display_df = pd.DataFrame({
        'Field': product_info.index,
        'Value': product_info.values
    })
    # Filter out empty/not specified values for cleaner display
    display_df = display_df[display_df['Value'].astype(str).str.strip() != '']
    display_df = display_df[display_df['Value'] != 'Not specified']
    
    st.table(display_df)
    
    st.markdown("---")
    
    # SEO Keywords Auto-Generation Section
    st.subheader("🔍 SEO Keywords for Amazon Content")
    
    # Initialize session state for SEO keywords
    if 'seo_keywords_data' not in st.session_state:
        st.session_state['seo_keywords_data'] = []
    if 'seo_keywords_text' not in st.session_state:
        st.session_state['seo_keywords_text'] = ""
    
    # Auto-generate button
    col1, col2 = st.columns([1, 4])
    with col1:
        if st.button("✨ Auto-Generate Keywords"):
            with st.spinner("🤖 Generating SEO Keywords..."):
                product_name = product_info.get('Product Name', 'N/A')
                key_benefits = product_info.get('USPs Front', 'N/A')
                description_summary = product_info.get('Claims', 'N/A')
                brand_tonality = BRAND_TONALITIES.get(selected_brand, '')
                
                keywords_data = generate_seo_keywords(product_name, key_benefits, description_summary, brand_tonality)
                st.session_state['seo_keywords_data'] = keywords_data
                
                # Convert to comma-separated text for easy editing
                if keywords_data:
                    keyword_list = [item.get('keyword', '') for item in keywords_data if item.get('keyword')]
                    st.session_state['seo_keywords_text'] = ", ".join(keyword_list)
                    st.success(f"✅ Generated {len(keywords_data)} SEO Keywords!")
    
    # Editable text area for keywords
    st.session_state['seo_keywords_text'] = st.text_area(
        "Edit SEO Keywords (comma-separated)",
        value=st.session_state['seo_keywords_text'],
        placeholder="e.g., anti-aging serum, wrinkle reducer, vitamin C, face serum for women",
        height=100,
        help="These keywords will be used in Amazon Title, Bullets, and Description. You can edit them before generating content."
    )
    
    # Clean and format keywords
    formatted_keywords = st.session_state['seo_keywords_text'].strip() if st.session_state['seo_keywords_text'] else ""
    
    if formatted_keywords:
        keyword_count = len([k.strip() for k in formatted_keywords.split(',') if k.strip()])
        st.info(f"📌 Using {keyword_count} keywords for Amazon content generation")
    else:
        st.warning("⚠️ Please generate or add SEO keywords before creating Amazon content")
    
    st.markdown("---")

    def generate_section(section_name, instruction):
        try:
            # Prepare the system message with brand tonality if available
            system_message = "You are a professional ecommerce copywriter."
            
            if selected_brand and selected_brand in BRAND_TONALITIES:
                brand_tonality = BRAND_TONALITIES[selected_brand]
                system_message += f"\n\nIMPORTANT - BRAND TONE & VOICE:\nYou are writing for the brand '{selected_brand}'. Follow this brand tonality strictly:\n\n{brand_tonality}\n\nEnsure all content reflects this brand's personality, tone, and style."
            
            completion = openai.chat.completions.create(
                model="gpt-4o-mini",
                messages=[
                    {"role": "system", "content": system_message},
                    {"role": "user", "content": instruction}
                ],
                temperature=0.7
            )
            return completion.choices[0].message.content.strip()
        except Exception as e:
            st.error(f"Error generating {section_name}: {e}")
            return ""

    def text_box(label, key, height):
        if key in st.session_state:
            st.text_area(label, st.session_state[key], height=height)

    # Amazon Content Generation
    st.subheader("🛒 Amazon Content")
    
    # Check if SEO keywords are present before allowing Amazon content generation
    amazon_disabled = not formatted_keywords
    
    if amazon_disabled:
        st.error("❌ SEO Keywords are required for Amazon content generation!")
        st.info("👆 Please generate or add SEO keywords in the section above before proceeding.")
    
    # Normalize brand name once for all Amazon content sections
    brand_name_raw = product_info.get('Brand Name', '').strip()
    normalized_brand = normalize_brand_name(brand_name_raw) or selected_brand or product_info.get('Brand Name', 'N/A')
    
    st.session_state.setdefault('title', '')
    if st.button("Generate Product Title", disabled=amazon_disabled):
        # Build SEO keywords instruction based on user input
        seo_instruction = ""
        if formatted_keywords:
            seo_instruction = f"\n- MUST include these SEO keywords naturally: {formatted_keywords}"
        else:
            seo_instruction = "\n- Add SEO keywords with high search volume & relevant keywords"
        
        # Check product category for category-specific title requirements
        category = product_info.get('Category', 'Beauty')
        is_electronics = 'Electronics' in str(category)
        
        if is_electronics:
            # Electronics-specific title requirements
            category_instruction = f"""
            
            ELECTRONICS CATEGORY - MUST INCLUDE IN TITLE:
            - Charging cable type (e.g., USB-C, Magnetic, Type-C)
            - Battery information (e.g., battery capacity, runtime, rechargeable)
            - Technology used (e.g., IPX7 waterproof, 5-speed motor, sonic technology)
            - These technical specifications are CRITICAL for electronics products
            """
        else:
            # Beauty category - net quantity requirement
            category_instruction = f"""
            
            BEAUTY CATEGORY - MUST INCLUDE IN TITLE:
            - Net Qty. in title with proper units (ml or g or units)
            - Example: "50ml", "30g", "60 Patches"
            - Net Quantity: {product_info.get('Net Qty.', 'N/A')}
            """
        
        title_prompt = f"""
            Write an Amazon product title (min 230 characters) for a high-converting listing.
            Use this format: Brand + Product Type + Keywords + Claims.
            Do not include size, weight, or volume dimensions in the title.
            
            IMPORTANT GUIDELINES:
            - Title character limit: 230-240 characters minimum
            - Do NOT use words like "bestselling" in the title
            - Do NOT use claims such as "whitening/brightening/fair" in the title
            - Instead, use words/phrases such as "reduces/promotes/helps in reducing"{category_instruction}{seo_instruction}
            - Ensure proper spelling and grammar

            Product: {product_info.get('Product Name', 'N/A')}
            Brand: {normalized_brand}
            Category: {category}
            USPs: {product_info.get('USPs Front', 'N/A')}
            Ingredients: {product_info.get('Ingredients', 'N/A')}
            Claims: {product_info.get('Claims', 'N/A')}
            Net Qty: {product_info.get('Net Qty.', 'N/A')}
            Know Your Product: {product_info.get('KNOW YOUR PRODUCT', '')}
        """
        st.session_state['title'] = generate_section("title", title_prompt)
    text_box("Amazon Product Title", 'title', 100)
    
    st.session_state.setdefault('bullets', '')
    if st.button("Generate Bullet Points", disabled=amazon_disabled):
        # Build SEO keywords instruction based on user input
        seo_instruction = ""
        if formatted_keywords:
            seo_instruction = f"\n- MUST include these SEO keywords naturally across the bullets: {formatted_keywords}"
        else:
            seo_instruction = "\n- Add SEO keywords with high search volume & relevant keywords"
        
        bullet_prompt = f"""
            Write 7 optimized Amazon bullet points. Each bullet should follow this format:
            BENEFIT IN CAPS: Followed by a clear, compelling benefit (250–300 characters).

            Avoid mentioning price, value for money, discounts, or any numerical cost-related details. Focus only on features, results, usage, or ingredient benefits.
            
            IMPORTANT GUIDELINES:
            - Do NOT use words like "bestselling"
            - Do NOT use claims such as "whitening/brightening/fair"
            - Instead, use words/phrases such as "reduces/promotes/helps in reducing"{seo_instruction}
            - Ensure proper spelling and grammar

            Product: {product_info.get('Product Name', 'N/A')}
            Brand: {normalized_brand}
            USPs: {product_info.get('USPs Front', 'N/A')}, {product_info.get('USP Back', 'N/A')}, {product_info.get('USP Side', 'N/A')}
            Ingredients: {product_info.get('Ingredients', 'N/A')}
            Claims: {product_info.get('Claims', 'N/A')}
            How to Use: {product_info.get('How to use it?', 'N/A')}
        """
        st.session_state['bullets'] = generate_section("bullet points", bullet_prompt)
    text_box("Generated Bullet Points", 'bullets', 200)

    st.session_state.setdefault('description', '')
    if st.button("Generate Description", disabled=amazon_disabled):
        # Build SEO keywords instruction based on user input
        seo_instruction = ""
        if formatted_keywords:
            seo_instruction = f"\n- MUST include these SEO keywords naturally in the description: {formatted_keywords}"
        else:
            seo_instruction = "\n- Add SEO keywords with high search volume & relevant keywords"
        
        desc_prompt = f"""
            Write an Amazon HTML product description (max 400 words).
            Use 2 paragraphs, <p> and <br> tags for light formatting.
            
            IMPORTANT GUIDELINES:
            - Do NOT use words like "bestselling"
            - Do NOT use claims such as "whitening/brightening/fair"
            - Instead, use words/phrases such as "reduces/promotes/helps in reducing"{seo_instruction}
            - Ensure proper spelling and grammar
            - Strictly avoid exaggerated exclamations/expressions in the description

            Product: {product_info.get('Product Name', 'N/A')}
            Brand: {normalized_brand}
            USPs: {product_info.get('USPs Front', 'N/A')}
            Ingredients: {product_info.get('Ingredients', 'N/A')}
            Claims: {product_info.get('Claims', 'N/A')}
            How to Use: {product_info.get('How to use it?', 'N/A')}
        """
        st.session_state['description'] = generate_section("description", desc_prompt)
    text_box("HTML Description", 'description', 150)

    st.session_state.setdefault('shopify', '')
    if st.button("Generate Shopify Description"):
        shopify_prompt = f"""
            Write a Shopify product description using a friendly, informative tone.
            Include small paragraphs, bullet points, and headings. Length: 1500–2000 characters.
            
            IMPORTANT: Do NOT use markdown formatting. Use plain text only.
            - Do not use ** for bold
            - Do not use # for headings
            - Do not use markdown syntax
            - Use simple text formatting only
            
            IMPORTANT: Strictly avoid exaggerated exclamations/expressions in the description.

            Product: {product_info.get('Product Name', 'N/A')}
            Brand: {normalized_brand}
            USPs: {product_info.get('USPs Front', 'N/A')}
            Ingredients: {product_info.get('Ingredients', 'N/A')}
            Claims: {product_info.get('Claims', 'N/A')}
            How to Use: {product_info.get('How to use it?', 'N/A')}
            MRP: {product_info.get('MRP (Incl. of all taxes)', 'N/A')}
        """
        st.session_state['shopify'] = generate_section("Shopify description", shopify_prompt)
    text_box("Shopify Description", 'shopify', 200)

    st.session_state.setdefault('hero', '')
    if st.button("Generate Hero Prompts"):
        # Use the selected category from the UI
        category = product_info.get('Category', 'Beauty')
        brand_name_raw = product_info.get('Brand Name', '').strip()
        brand_name = normalize_brand_name(brand_name_raw) or selected_brand
        
        # Brand-specific tone instruction
        brand_tone = ""
        if brand_name and brand_name in BRAND_TONALITIES:
            brand_tone = f"\n\nBRAND TONE: {brand_name}\n{BRAND_TONALITIES[brand_name]}\nEnsure all headlines and text overlays reflect this brand's voice and personality."
        
        if 'Beauty' in str(category):
            hero_prompt = f"""
Generate 8 hero image prompts for a beauty product listing (1500 x 1500 pixels each).
IMPORTANT: Provide the output in plain text format only. Do NOT use markdown syntax like **, ##, or bullet points with *. Use simple numbered lists and plain text.

CRITICAL NOTE: Hero and A+ images should NOT repeat the same content or visuals. Even if the core message aligns, it must be expressed differently through distinct headlines, layouts, and imagery.{brand_tone}

Hero Image 1: Product Highlight
Content Focus: Always provide these 3 variations
Hero Image 1.a_Product Image (clean product shot on neutral background)
Hero Image 1.b_Product + Box (product with packaging)
Hero Image 1.c_Product + Model (product with model showcasing usage or application)

Hero Image 2: Core Value Proposition
Content Focus: Show what the product does best and the main problem it solves.
Examples: "Hydration that Lasts 72 Hours", "Brightens, Firms, and Smoothens", "Glass Skin Made Simple"
Goal: Hook emotionally + highlight the transformation
Visual Direction: Model or texture-based visual (glowing skin close-up, droplets). Use warm lighting and a big, bold headline with minimal supporting text.

Hero Image 3: Benefits Snapshot
Content Focus: List 3-5 clear user-facing benefits based on USPs and ingredients
Examples: Deep Hydration, Reduces Dullness, Refines Pores, Enhances Glow, Lightweight Formula
Goal: Educate quickly about visible results customers can expect
Visual Direction: Icon grid or short text blocks next to corresponding visuals (skin texture, serum drop, flower ingredient, glow effect). Use clean brand colors and symmetrical spacing.

Hero Image 4: Ingredient Power / Science Behind
Content Focus: Highlight key actives or formulation technology
Examples: "Powered by 5% Niacinamide + Hyaluronic Acid", "Infused with Rosehip & Avocado Oils", "K-Beauty Formula for Glass-Like Skin"
Goal: Show authenticity and formulation strength
Visual Direction: Macro ingredient visuals (bubbles, botanicals, lab glass textures). Combine text overlays and ingredient icons. Use light backgrounds to emphasize purity.

Hero Image 5: Comparison Table / Why It's Better
Content Focus: Show how your product outperforms typical alternatives (creams, masks, etc.)
Goal: Build trust and show superiority
Visual Direction: Clean two-column table with brand-color highlights. Add product image near your column.

Hero Image 6: How to Use / Routine Step
Content Focus: Simple, 3-4 step usage guide
Example: 1. Cleanse your skin, 2. Apply a thin layer evenly, 3. Leave overnight or rinse as directed, 4. Wake up to soft, radiant skin
Goal: Simplify usage and make it look effortless
Visual Direction: Step-by-step icons or close-up lifestyle visuals of hand applying product. Use numbered circles and pastel backgrounds.

Hero Image 7: Key Differentiator / Main USP
Content Focus: Highlight one powerful claim or innovation from USPs and Ingredient list
Examples: "50,000 PPM Niacinamide - The Glass Skin Secret", "Clinically Proven to Hydrate Overnight", "Formulated by Korean Skin Experts"
Goal: Create a memorable "hero claim"
Visual Direction: Close-up of texture or ingredient with overlay text. Add scientific/clean design cues (beaker glow, droplet macro).

Hero Image 8: Before and After using the Product
Content Focus: Visually showcase the transformation - smoother, clearer, or more radiant skin
Goal: Build trust by showing tangible, visible improvements
Visual Direction: Side-by-side layout ("Before" left, "After" right). Use consistent lighting and angles. Add soft labels and minimal caption like "Brighter. Smoother. More Hydrated." Include product image subtly near "After" side. Clean, bright, natural background.

Product details:
Product: {product_info.get('Product Name', 'N/A')}
USPs: {product_info.get('USPs Front', 'N/A')}
Ingredients: {product_info.get('Ingredients', 'N/A')}
Claims: {product_info.get('Claims', 'N/A')}
How to use: {product_info.get('How to use it?', 'N/A')}
Target: {product_info.get('Target Audience', '')}
"""
        elif 'Electronics' in str(category):
            hero_prompt = f"""
Generate 8 hero image prompts for an electronics product listing (1500 x 1500 pixels each).
IMPORTANT: Provide the output in plain text format only. Do NOT use markdown syntax like **, ##, or bullet points with *. Use simple numbered lists and plain text.

CRITICAL NOTE: Hero and A+ images should NOT repeat the same content or visuals. Even if the core message aligns, it must be expressed differently through distinct headlines, layouts, and imagery.{brand_tone}

Hero Image 1: Product Highlight
Content Focus: Always provide these 3 variations
Hero Image 1.a_Product Image (clean product shot on neutral background)
Hero Image 1.b_Product + Box (product with packaging)
Hero Image 1.c_Product + Model (product being used by a model)

Hero Image 2: Core Value Proposition
Content Focus: Address what the product does, key pain points or highlight main benefit
Examples: "Save thousands on Salon Bills", "Edge Shave - Zero Cuts. Chiseled Jaw Line", "Meet Rizzler- the beard boss"
Goal: Hook users emotionally or practically within 3 seconds
Visual Direction: Product-in-use shot showing a real person. Use short, bold headline + smaller subtext below. Include supporting visual cue (glowing motor light, water resistance, before-after contrast).

Hero Image 3: Key Areas of Use / Versatility
Content Focus: Show how and where the product can be used
Examples: "Face • Body • Underarms • Bikini" for trimmers, "Blow • Curl • Straighten" for hair tools
Goal: Demonstrate multi-use versatility
Visual Direction: Use product with labeled zones or side-by-side images. Add small icons or silhouettes for clarity. Keep layout clean and symmetrical.

Hero Image 4: Feature Highlights / Performance Benefits
Content Focus: Show top 4-6 standout features in short, clear bullets
Examples: 3 Interchangeable Heads, Dual-Speed Motor, IPX6 Waterproof, Type-C Charging, 100 Mins Runtime, Skin-Safe Blades
Goal: Summarize core specs + key benefits in one glance
Visual Direction: Split layout with product in center, feature callouts surrounding it. Use iconography or minimal graphic arrows to highlight components.

Hero Image 5: Comparison Table
Content Focus: Show why your product is better than others - without naming competitors. Include 5-6 comparison points in tabular format
Goal: Prove technical and functional superiority
Visual Direction: Clean, minimal table design with brand-color highlights. Keep text large and legible; use product image or icon beside your column.

Hero Image 6: What's in the Box
Content Focus: Show exactly what the customer receives. Only list all components (main device, attachments, charger, brush, manual, etc.)
Goal: Set clear buyer expectations and highlight value
Visual Direction: Flat lay or 3D arrangement on neutral background. Use labels or numbering for each item. Ensure lighting consistency and neat spacing.

Hero Image 7: Key Differentiator / Main USP
Content Focus: Highlight one standout USP or technology that defines your product
Examples: "Five heads and where they can be used", "Ceramic Blades for Safe Precision", "Dual-Speed Motor - Power Meets Control"
Goal: Make the customer remember the key innovation or value
Visual Direction: Close-up macro shot or bold visual metaphor (charging cable glowing, airflow effect, blade detail). Big product name overlay with short USP line underneath.

Hero Image 8: Lifestyle / Versatility Slide
Content Focus: Show different looks or outcomes achievable with one tool
Examples: "One Tool. Different Looks", "From Office to Party - Styled in Minutes", "Your Everyday Styling Partner"
Goal: Highlight creative versatility and emotional payoff
Visual Direction: Collage or side-by-side layout showing multiple styles or use-cases (beard shapes, hair looks). Keep tone aspirational and modern.

Product details:
Product: {product_info.get('Product Name', 'N/A')}
Main USP: {product_info.get('USPs Front', 'N/A')}
Other USPs: {product_info.get('USP Back', '')}
Box Includes: {product_info.get('Box Includes', 'N/A')}
How to use: {product_info.get('How to use it?', 'N/A')}
Other details: {product_info.get('KNOW YOUR PRODUCT', '')}
"""
        else:
            hero_prompt = "Category not supported. Please specify 'Beauty' or 'Electronics' in the Category field."

        st.session_state['hero'] = generate_section("hero image prompts", hero_prompt)
    text_box("Hero Image Prompts", 'hero', 180)

    st.session_state.setdefault('a_plus', '')
    if st.button("Generate A+ Prompts"):
        # Get category and brand
        category = product_info.get('Category', 'Beauty')
        brand_name_raw = product_info.get('Brand Name', '').strip()
        brand_name = normalize_brand_name(brand_name_raw) or selected_brand
        
        # Brand-specific tone instruction
        brand_tone = ""
        if brand_name and brand_name in BRAND_TONALITIES:
            brand_tone = f"\n\nBRAND TONE: {brand_name}\n{BRAND_TONALITIES[brand_name]}\nEnsure all headlines and text overlays reflect this brand's voice and personality."
        
        if 'Beauty' in str(category):
            a_plus_prompt = f"""
Generate 6 Amazon A+ content image prompts for a beauty product.
IMPORTANT: Provide the output in plain text format only. Do NOT use markdown syntax like **, ##, or bullet points with *. Use simple numbered lists and plain text.

CRITICAL NOTE: A+ images should expand the story by covering features, benefits, or brand details that are NOT highlighted in the hero images. Do not repeat hero image content.{brand_tone}

A+ Image 1: Hero Banner + Core Promise
Content Focus: Introduce the product and its biggest selling point in one frame
- Short headline (6-8 words) that captures what the product does best
- One-line subtext describing the benefit or problem it solves
Example: Headline "Trim Smart. Style Effortlessly." | Subtext "Advanced precision grooming with zero nicks or cuts."
Goal: Instantly communicate product purpose + emotional appeal (confidence, convenience, comfort)
Visual Direction: Big, clean hero shot with product angled toward camera, soft gradient or lifestyle background (bathroom/vanity setup), short text overlay. Use a person in frame for relatability.

A+ Image 2: Key Benefits & Features Overview
Content Focus: Summarize top 4-6 performance benefits. Each point short, clear, action-oriented
Goal: Show at a glance what makes the product technically strong and convenient
Visual Direction: Grid layout or circular icons with micro close-ups of each feature. Clean background with product in center.

A+ Image 3: Ingredients & Science
Content Focus: Show hero ingredients and their benefits
Examples: Niacinamide - Brightens and evens tone, Hyaluronic Acid - Deep hydration, Rosehip Oil - Restores softness
Goal: Build trust and education
Visual Direction: Macro ingredient visuals (drops, botanicals, molecular patterns), short text blocks, clean pastel background.

A+ Image 4: How to Use (Pre & Post Steps)
Content Focus: Show complete skincare ritual - Before Use (clean, dry skin), During Use (apply evenly with upward strokes), After Use (let absorb or rinse), Post Care (follow with moisturizer/sunscreen)
Goal: Make application steps easy to follow while positioning as self-care ritual
Visual Direction: Soft, minimal design with numbered steps. Calm pastel tones, water or glow textures, close-up visuals of model or hand applying product.

A+ Image 5: Other Products from the Same Brand
Content Focus: Showcase full skincare routine from same brand to encourage brand loyalty and cross-selling
Example caption: "Complete Your Glow Routine with {brand_name}"
Goal: Build brand ecosystem and communicate that products work better together
Visual Direction: Flat lay or lineup of all products on marble or pastel background.

A+ Image 6: Customer Reviews & FAQs
Content Focus: Highlight authentic customer feedback that subtly answers common skincare doubts
Sample Reviews with star ratings addressing: texture & results, safety for sensitive skin, frequency & visible benefits, absorption & effect
Goal: Build social proof and credibility while resolving common queries
Visual Direction: Soft review cards with real-sounding Indian names and 5-star icons.

Product details:
Product: {product_info.get('Product Name', 'N/A')}
USPs: {product_info.get('USPs Front', 'N/A')}
Ingredients: {product_info.get('Ingredients', 'N/A')}
Claims: {product_info.get('Claims', 'N/A')}
How to Use: {product_info.get('How to use it?', 'N/A')}
"""
        elif 'Electronics' in str(category):
            a_plus_prompt = f"""
Generate 6 Amazon A+ content image prompts for an electronics product.
IMPORTANT: Provide the output in plain text format only. Do NOT use markdown syntax like **, ##, or bullet points with *. Use simple numbered lists and plain text.

CRITICAL NOTE: A+ images should expand the story by covering features, benefits, or brand details that are NOT highlighted in the hero images. Do not repeat hero image content.{brand_tone}

A+ Image 1: Hero Banner + Core Promise
Content Focus: Introduce the product and its biggest selling point in one frame
- Short headline (6-8 words) that captures what the product does best
- One-line subtext describing the benefit or problem it solves
Example: Headline "Trim Smart. Style Effortlessly." | Subtext "Advanced precision grooming with zero nicks or cuts."
Goal: Instantly communicate product purpose + emotional appeal (confidence, convenience, comfort)
Visual Direction: Big, clean hero shot with product angled toward camera, soft gradient or lifestyle background (bathroom/vanity setup), short text overlay. Use a person in frame for relatability.

A+ Image 2: Key Benefits & Features Overview
Content Focus: Summarize top 4-6 performance benefits. Each point short, clear, action-oriented
Examples: Dual-Speed Motor for Precision, Type-C Fast Charging, IPX6 Waterproof, 100 Mins Runtime, Skin-Safe Ceramic Blades, Low-Noise Operation
Goal: Show at a glance what makes the product technically strong and convenient
Visual Direction: Grid layout or circular icons with micro close-ups of each feature (motor, blades, water resistance, battery). Clean background with product in center.

A+ Image 3: How to Use - Step-by-Step
Content Focus: Simplify the process visually with 3-4 short steps: Attach the right head/setting, Power on and glide gently, Clean and rinse after use, Store safely for next use
Goal: Make it look effortless and beginner-friendly
Visual Direction: Show real hands using the product. Use numbered steps with minimal text and small arrows/icons for direction. Keep color palette light, text clean and readable.

A+ Image 4: Technology & Safety
Content Focus: Highlight tech that ensures performance and user safety
Example points: Smart Motor Control, Overheat Protection/Auto Shut-off, Ergonomic grip, Hypoallergenic material or skin-safe design
Goal: Build trust in quality, durability, and innovation
Visual Direction: Split image layout - left side showing macro or exploded view of product components (motor, battery, heating system), right side showing clean bullet list or icon-based text.

A+ Image 5: Easy Charging & Travel-Friendly
Content Focus: Highlight portability and convenience
Include: Type-C Universal Charging (charge via laptop, adapter, power bank), Long-lasting Battery (runtime on one charge), Compact & Lightweight (fits in travel pouch)
Goal: Reinforce "carry it anywhere, use it anywhere" confidence
Visual Direction: Product next to laptop, power bank, or travel pouch. Include icons/graphics showing charging cable connection. Soft gradient or lifestyle setup (dresser, hotel counter).

A+ Image 6: Promoting different products of the same brand
Content Focus: Promote other products from same brand to encourage cross-selling. Include 3-4 product visuals from related categories
Example: "Try our Fan-Favorites" or "Choose from the range of other products"
Goal: Build brand recall and drive multi-product purchase
Visual Direction: Clean product grid or carousel layout with uniform lighting.

Product details:
Product: {product_info.get('Product Name', 'N/A')}
Main USP: {product_info.get('USPs Front', 'N/A')}
Other USPs: {product_info.get('USP Back', '')}
Box Includes: {product_info.get('Box Includes', 'N/A')}
How to use: {product_info.get('How to use it?', 'N/A')}
Battery/Charging: {product_info.get('Battery', 'N/A')}
"""
        else:
            a_plus_prompt = "Category not supported. Please specify 'Beauty' or 'Electronics' in the Category field."
        
        st.session_state['a_plus'] = generate_section("A+ image prompts", a_plus_prompt)
    text_box("A+ Image Prompts", 'a_plus', 180)

    # Website Content Generation
    if st.button("🌐 Generate Full Website Content"):
        # Use the selected category from the UI
        category = product_info.get('Category', 'Beauty')
        is_electronics = 'Electronics' in str(category)
        
        # Normalize brand name for consistent detection
        brand_name_raw = product_info.get('Brand Name', '').strip()
        brand_name = normalize_brand_name(brand_name_raw) or selected_brand
        
        # Check brand using normalized name
        is_seoulskin = brand_name == 'Seoulskin'
        is_urban_gabru = brand_name == 'Urban Gabru'
        is_urban_yog = brand_name == 'Urban Yog'
        is_makemeebold = brand_name == 'MakeMeeBold'
        
        # Check if product is unisex from 'Ideal For (Gender)' field
        gender_field = product_info.get('Ideal For (Gender)', '').lower()
        is_unisex = 'unisex' in gender_field or ('men' in gender_field and 'women' in gender_field) or ('male' in gender_field and 'female' in gender_field)
        
        # Determine language distribution based on brand
        if is_seoulskin:
            if is_unisex:
                review_language_instruction = """15 Customer Reviews - SEOULSKIN (PREMIUM ENGLISH FOCUS):
               - MAJORITY (12 reviews) should be in premium English with sophisticated tone
               - MINORITY (3 reviews) can include light Hinglish for authenticity
               - English reviews should emphasize K-beauty, innovation, luxury, premium skincare
               - Keep Hinglish reviews minimal and subtle (e.g., "bohot acha", "mujhe pasand aaya")
               - Focus on premium vocabulary: "radiant", "luminous", "transformative", "innovative"
               - GENDER DISTRIBUTION (UNISEX): 12 female reviews + 2 male reviews
               - Include 1 review mentioning product as a gifting option
               - Include comparisons like "better than [old product/technology]"
               - Mention relatable problems: "This helps because I have [common problem]"
               - Use common Indian names (first + last name)
               - Format: Name – Review (no markdown, just plain text)"""
            else:
                review_language_instruction = """15 Customer Reviews - SEOULSKIN (PREMIUM ENGLISH FOCUS):
               - MAJORITY (10-12 reviews) should be in premium English with sophisticated tone
               - MINORITY (3-5 reviews) can include light Hinglish for authenticity
               - English reviews should emphasize K-beauty, innovation, luxury, premium skincare
               - Keep Hinglish reviews minimal and subtle (e.g., "bohot acha", "mujhe pasand aaya")
               - Focus on premium vocabulary: "radiant", "luminous", "transformative", "innovative"
               - Include 1 review mentioning product as a gifting option
               - Include comparisons like "better than [old product/technology]"
               - Mention relatable problems: "This helps because I have [common problem]"
               - Use common Indian names (first + last name)
               - Format: Name – Review (no markdown, just plain text)"""
        elif is_urban_yog:
            if is_unisex:
                review_language_instruction = """15 Customer Reviews - URBAN YOG:
               - 10 reviews in Hinglish with some Gen Z words but simple
               - 5 reviews in English (premium tone)
               - GENDER DISTRIBUTION (UNISEX): 13 female reviews + 2 male reviews
               - Include 1 review mentioning product as a gifting option
               - Include comparisons like "better than [old product/technology]"
               - Mention relatable problems: "This helps because I have [common problem]"
               - Use common Indian names (first + last name)
               - Format: Name – Review (no markdown, just plain text)"""
            else:
                review_language_instruction = """15 Customer Reviews - URBAN YOG:
               - 10 reviews in Hinglish with some Gen Z words but simple
               - 5 reviews in English (premium tone)
               - Include 1 review mentioning product as a gifting option
               - Include comparisons like "better than [old product/technology]"
               - Mention relatable problems: "This helps because I have [common problem]"
               - Use common Indian names (first + last name)
               - Format: Name – Review (no markdown, just plain text)"""
        elif is_makemeebold:
            if is_unisex:
                review_language_instruction = """15 Customer Reviews - MAKEMEEBOLD:
               - 10 reviews in English but in simple tone
               - 5 reviews in Hinglish with English words
               - GENDER DISTRIBUTION (UNISEX): 13 female reviews + 2 male reviews
               - Include 1 review mentioning product as a gifting option
               - Include comparisons like "better than [old product/technology]"
               - Mention relatable problems: "This helps because I have [common problem]"
               - Use common Indian names (first + last name)
               - Format: Name – Review (no markdown, just plain text)"""
            else:
                review_language_instruction = """15 Customer Reviews - MAKEMEEBOLD:
               - 10 reviews in English but in simple tone
               - 5 reviews in Hinglish with English words
               - Include 1 review mentioning product as a gifting option
               - Include comparisons like "better than [old product/technology]"
               - Mention relatable problems: "This helps because I have [common problem]"
               - Use common Indian names (first + last name)
               - Format: Name – Review (no markdown, just plain text)"""
        elif is_urban_gabru:
            # Urban Gabru is men's brand - no unisex consideration
            review_language_instruction = """15 Customer Reviews - URBAN GABRU (MEN'S BRAND):
               - 10 reviews in Hinglish and simple
               - 5 reviews in English
               - All reviews from male customers (use male names)
               - Include 1 review mentioning product as a gifting option
               - Include comparisons like "better than [old product/technology]"
               - Mention relatable problems: "This helps because I have [common problem]"
               - Use common Indian names (first + last name)
               - Format: Name – Review (no markdown, just plain text)"""
        else:
            review_language_instruction = """15 Customer Reviews in HINGLISH format:
               - Use natural Hinglish (Hindi-English mix) like: "bohot acha h", "mujhe pasand aaya", "ekdum fresh lagta h"
               - First 2 reviews: Long story-style testimonials (premium English with Hindi words)
               - Remaining 13 reviews: Short casual Hinglish reviews
               - Mix casual Hinglish with some premium English reviews
               - Include 1 review mentioning product as a gifting option
               - Include comparisons like "better than [old product/technology]"
               - Mention relatable problems: "This helps because I have [common problem]"
               - Use common Indian names (first + last name)
               - Format: Name – Review (no markdown, just plain text)"""
        
        if is_electronics:
            # Electronics-specific full content with Box Includes and Warranty
            full_web_prompt = f"""
            Generate the following website content in a structured format:

            1. 7 Bullet Points – focus on customer benefits, pain points, or product uniqueness.
            2. Description – 2 paragraphs, warm and benefit-driven tone, max 400 words.
            3. USP Points – concise value-driven phrases (2–4 words each).
            4. What's in the Box – List items with format: -[Quantity] [Item Name] ([Details if applicable]), keep simple, do not exaggerate.
            5. Warranty Information – Include warranty details clearly.
            6. How to use – easy-to-follow, friendly, step-by-step instructions.
            7. 6 FAQs – ELECTRONICS FORMAT:
               Focus on practical questions that genuinely help customers:
               - Usage limitations (e.g., "Can I use it on my face?")
               - Water resistance/waterproof rating
               - Battery life and charging time
               - What's included in the package
               - Suitability for different needs
               - Warranty information
               FORMAT: Q1: [Question] / A: [Concise answer with specific details like ratings, numbers, specifications]
               Keep answers under 150 characters and be specific with technical details.
            8. {review_language_instruction}
            9. Brand & Contact Information – Include brand ownership, country of origin, contact details.
            
            IMPORTANT: Provide ALL content in plain text format only. Do NOT use markdown formatting such as:
            - No ** for bold text
            - No ## for headings
            - No *** or ___ for dividers
            - No markdown bullet points with *
            - Use simple text formatting only with clear section labels
            
            IMPORTANT: Strictly avoid exaggerated exclamations/expressions in the description.

            Product Name: {product_info.get('Product Name', 'N/A')}
            Brand: {brand_name}
            USPs / Features: {product_info.get('USPs Front', 'N/A')}
            Ingredients: {product_info.get('Ingredients', 'N/A')}
            Claims: {product_info.get('Claims', 'N/A')}
            How to Use: {product_info.get('How to use it?', 'N/A')}
            MRP: {product_info.get('MRP (Incl. of all taxes)', 'N/A')}
            Brand Owned & Marketed By: {product_info.get('Brand Owned & Marketed By', 'N/A')}
            Country of Origin: {product_info.get('Country Of Origin', 'N/A')}
            Email: {product_info.get('Email', 'N/A')}
            Contact: {product_info.get('Contact Us', 'N/A')}
            Box Includes: {product_info.get('Box Includes', 'N/A')}
            Warranty: {product_info.get('Warranty', 'N/A')}
        """
        else:
            # Beauty products - with "How Much Do You Get?" instead of "What's in the Box?"
            full_web_prompt = f"""
            Generate the following website content in a structured format:

            1. 7 Bullet Points – focus on customer benefits, pain points, or product uniqueness.
            2. Description – 2 paragraphs, warm and benefit-driven tone, max 400 words.
            3. USP Points – concise value-driven phrases (2–4 words each).
            4. How Much Do You Get? – Describe the net quantity/volume (ml, g, units, patches, etc.) clearly and simply.
            5. How to use – easy-to-follow, friendly, step-by-step instructions.
            6. 6 FAQs – with short, helpful answers from user perspectives.
            7. {review_language_instruction}
            8. Brand & Contact Information – Include brand ownership, country of origin, contact details.
            
            IMPORTANT: Provide ALL content in plain text format only. Do NOT use markdown formatting such as:
            - No ** for bold text
            - No ## for headings
            - No *** or ___ for dividers
            - No markdown bullet points with *
            - Use simple text formatting only with clear section labels
            
            IMPORTANT: Strictly avoid exaggerated exclamations/expressions in the description.

            Product Name: {product_info.get('Product Name', 'N/A')}
            Brand: {brand_name}
            USPs / Features: {product_info.get('USPs Front', 'N/A')}
            Ingredients: {product_info.get('Ingredients', 'N/A')}
            Claims: {product_info.get('Claims', 'N/A')}
            How to Use: {product_info.get('How to use it?', 'N/A')}
            Net Quantity/Weight: {product_info.get('Net Qty.', 'N/A')}
            MRP: {product_info.get('MRP (Incl. of all taxes)', 'N/A')}
            Brand Owned & Marketed By: {product_info.get('Brand Owned & Marketed By', 'N/A')}
            Country of Origin: {product_info.get('Country Of Origin', 'N/A')}
            Email: {product_info.get('Email', 'N/A')}
            Contact: {product_info.get('Contact Us', 'N/A')}
        """
        
        web_full = generate_section("Website Content", full_web_prompt)
        st.session_state['web_full'] = web_full
    text_box("Full Website Content", 'web_full', 500)
    st.subheader("🌐 Website Content")

    st.session_state.setdefault('web_bullets', '')
    if st.button("Generate Website Bullet Points"):
        web_bullet_prompt = f"""
            Write 7 bullet points for website use. Each point should focus on product benefits, customer problems, or emotional triggers.

            Product: {product_info.get('Product Name', 'N/A')}
            USPs: {product_info.get('USPs Front', 'N/A')}
        """
        st.session_state['web_bullets'] = generate_section("website bullets", web_bullet_prompt)
    text_box("Website Bullet Points", 'web_bullets', 200)

    st.session_state.setdefault('web_description', '')
    if st.button("Generate Website Description"):
        # Use the selected category from the UI
        category = product_info.get('Category', 'Beauty')
        is_electronics = 'Electronics' in str(category)
        
        # Normalize brand name for consistency
        brand_name_raw = product_info.get('Brand Name', '').strip()
        normalized_brand_web = normalize_brand_name(brand_name_raw) or selected_brand or product_info.get('Brand Name', 'N/A')
        
        if is_electronics:
            # Electronics-specific description with Box Includes and Warranty
            web_desc_prompt = f"""
            Write a product description for the website (max 400 words). Use a warm, benefit-oriented tone and structure it with three paragraphs.
            Include brand credibility information naturally in the content.
            Incorporate relevant keywords and highlight the product's USPs throughout the description.
            
            After the main description, add two additional sections:
            1. "What's in the Box" - List all items included as bullet points (one item per line with a dash)
            2. "Warranty Information" - Include warranty details
            
            Format the box contents as simple bullet points (use - for each item, one per line).
            
            IMPORTANT: Strictly avoid exaggerated exclamations/expressions in the description.

            Product: {product_info.get('Product Name', 'N/A')}
            Brand: {normalized_brand_web}
            USPs: {product_info.get('USPs Front', 'N/A')}
            Claims: {product_info.get('Claims', 'N/A')}
            Ingredients: {product_info.get('Ingredients', 'N/A')}
            Brand Owned & Marketed By: {product_info.get('Brand Owned & Marketed By', 'N/A')}
            Country of Origin: {product_info.get('Country Of Origin', 'N/A')}
            Box Includes: {product_info.get('Box Includes', 'N/A')}
            Warranty: {product_info.get('Warranty', 'N/A')}
        """
        else:
            # Beauty products - include "How Much Do You Get?" section
            web_desc_prompt = f"""
            Write a product description for the website (max 400 words). Use a warm, benefit-oriented tone and structure it with three paragraphs.
            Include brand credibility information naturally in the content.
            Incorporate relevant keywords and highlight the product's USPs throughout the description.
            
            After the paragraphs, add a section titled "How Much Do You Get?" that describes the net quantity/volume included (e.g., "60 patches (30 pairs)", "50ml", "100g"). Use the Net Quantity/Weight field below.
            
            IMPORTANT: Strictly avoid exaggerated exclamations/expressions in the description.

            Product: {product_info.get('Product Name', 'N/A')}
            Brand: {normalized_brand_web}
            USPs: {product_info.get('USPs Front', 'N/A')}
            Claims: {product_info.get('Claims', 'N/A')}
            Ingredients: {product_info.get('Ingredients', 'N/A')}
            Net Quantity/Weight: {product_info.get('Net Qty.', 'N/A')}
            Brand Owned & Marketed By: {product_info.get('Brand Owned & Marketed By', 'N/A')}
            Country of Origin: {product_info.get('Country Of Origin', 'N/A')}
        """
        st.session_state['web_description'] = generate_section("website description", web_desc_prompt)
    text_box("Website Description", 'web_description', 200)

    st.session_state.setdefault('usp', '')
    if st.button("Generate USP"):
        usp_prompt = f"""
            Generate exactly 6 concise USPs for a product listing.
            Guidelines:
            - PRIORITY ORDERING: If there are any interchangeable heads or new type of technology, these MUST be placed on top (position 1)
            - Focus on the most important USPs first
            - Order USPs by priority: Most important → Least important
            - The first 2-3 USPs should be the strongest differentiators
            - Each USP must be a short, impactful phrase of 2–4 words.
            - Do not include full sentences.
            - Do not include explanations.
            - Format strictly as a bullet list like:
            - Frizz-control Formula
            - Lightweight & Non-Greasy
            - Safe for Daily Use

            Product: {product_info.get('Product Name', 'N/A')}
            USPs: {product_info.get('USPs Front', 'N/A')}
        """
        st.session_state['usp'] = generate_section("USP", usp_prompt)
    text_box("USP", 'usp', 100)

    st.session_state.setdefault('what_you_get', '')
    
    # Check category to determine which section to show
    category = product_info.get('Category', 'Beauty')
    is_electronics = 'Electronics' in str(category)
    
    if is_electronics:
        # Electronics: "What's in the Box?" section
        if st.button("Generate 'What's in the Box?'"):
            get_prompt = f"""
                Describe what the customer will receive when they purchase the product.
                
                IMPORTANT: Use the title "What's in the Box?" (not "What Do You Get?")
                
                FORMAT REQUIREMENTS:
                - List each item on a separate line starting with a dash (-)
                - Start each item with quantity number (e.g., -1, -5, -2)
                - Format: -[Quantity] [Item Name] ([Color/Detail if applicable]), keep simple, do not exaggerate.
                - Example format:
                  -1 Auto Electric Bath Brush (Off-White)
                  -5 Brush Heads: Long Bristle, Short Bristle, Silicone, Scrub, Gauze
                  -1 Magnetic USB Charging Cable
                  -1 Wall-Mount Hook with Adhesive Tape
                  -1 User Manual
                  -1 Warranty Card

                Product: {product_info.get('Product Name', 'N/A')}
                Net Quantity: {product_info.get('Net Qty.', 'N/A')}
                Box Includes: {product_info.get('Box Includes', 'N/A')}
            """
            st.session_state['what_you_get'] = generate_section("What's in the Box?", get_prompt)
        text_box("What's in the Box?", 'what_you_get', 100)
    else:
        # Beauty: "How Much Do You Get?" section
        if st.button("Generate 'How Much Do You Get?'"):
            get_prompt = f"""
                Describe the quantity/volume that the customer will receive when they purchase the product.
                
                IMPORTANT: Use the title "How Much Do You Get?" (not "What Do You Get?")
                
                FORMAT REQUIREMENTS:
                - Focus on the net quantity/weight/volume of the product
                - Be clear and specific about measurements (ml, g, units, patches, etc.)
                - Mention if it's a single unit or multiple units
                - Keep it simple and straightforward
                - Example formats:
                  "You get 60 hydrogel under-eye patches (30 pairs) in each jar."
                  "This bottle contains 50ml of serum for daily use."
                  "Each pack includes 100g of face cream."

                Product: {product_info.get('Product Name', 'N/A')}
                Net Quantity/Weight: {product_info.get('Net Qty.', 'N/A')}
            """
            st.session_state['what_you_get'] = generate_section("How Much Do You Get?", get_prompt)
        text_box("How Much Do You Get?", 'what_you_get', 100)

    st.session_state.setdefault('how_to_use', '')
    if st.button("Generate 'How to Use'"):
        # Check product category for category-specific format
        category = product_info.get('Category', 'Beauty')
        is_electronics = 'Electronics' in str(category)
        
        if is_electronics:
            # Electronics - use specific numbered format with device-focused instructions
            how_to_prompt = f"""
            Describe how to use this electronics product step-by-step in a friendly, instructional tone.
            
            FORMAT REQUIREMENTS (ELECTRONICS):
            - Use numbered steps (1, 2, 3, etc.)
            - Each step should be clear, concise, and actionable
            - Start each step with a verb (action word)
            - Keep each step to one sentence
            - Focus on device operation, power/charging, usage, maintenance
            - Example format:
              1. Select and attach your preferred brush head.
              2. Long press the power button to turn ON (choose speed as needed).
              3. Move in gentle circular motions on skin.
              4. Detach & rinse the brush head after each use.
              5. Dry & store on the wall-mount holder.
              6. Recharge using the magnetic 2-pin cable when required.

            Product: {product_info.get('Product Name', 'N/A')}
            Instructions: {product_info.get('How to use it?', 'N/A')}
            """
        else:
            # Beauty - use data from KLD sheet "How to use it?" column
            how_to_prompt = f"""
            Create clear, step-by-step usage instructions for this beauty product based on the provided instructions from the product sheet.
            
            FORMAT REQUIREMENTS (BEAUTY):
            - Use numbered steps (1, 2, 3, etc.)
            - Each step should be clear, concise, and actionable
            - Start each step with a verb (action word)
            - Keep each step to one sentence
            - Focus on application method, timing, frequency, and results
            - Base the content strictly on the "How to use it?" data provided below

            Product: {product_info.get('Product Name', 'N/A')}
            Instructions from sheet: {product_info.get('How to use it?', 'N/A')}
            """
        
        st.session_state['how_to_use'] = generate_section("How to Use", how_to_prompt)
    text_box("How to Use This", 'how_to_use', 100)

    st.session_state.setdefault('faqs', '')
    if st.button("Generate FAQs"):
        # Check product category for electronics-specific FAQ format
        category = product_info.get('Category', 'Beauty')
        is_electronics = 'Electronics' in str(category)
        
        if is_electronics:
            # Electronics-specific FAQ format with practical questions
            faqs_prompt = f"""
            Write 6 frequently asked questions and concise answers for this electronics product listing.
            
            IMPORTANT - ELECTRONICS FAQ FORMAT:
            Use this Q&A format with questions that genuinely help customers make informed decisions.
            
            Focus on these types of practical questions:
            - Usage limitations (e.g., "Can I use it on my face?")
            - Water resistance/waterproof rating
            - Battery life and charging time
            - What's included in the package
            - Suitability for different needs (e.g., sensitive skin, hair types)
            - Warranty information
            - Power specifications
            - Compatibility
            - Maintenance and care
            
            FORMAT:
            Q1: [Question]
            A: [Concise answer - be specific with details like ratings, numbers, specifications]
            
            EXAMPLE:
            Q1: Can I use it on my face?
            A: No, this brush is designed for body and back use only.
            Q2: Is it waterproof?
            A: Yes, it has an IPX7 waterproof rating. Safe for shower use but do not submerge the main unit.
            
            Guidelines:
            - Keep answers under 150 characters
            - Be specific with technical details (ratings, times, quantities)
            - Use clear, customer-friendly language
            - Do not mention price, promotions, or other brands

            Product: {product_info.get('Product Name', 'N/A')}
            Key Features: {product_info.get('USPs Front', 'N/A')}
            How to Use: {product_info.get('How to use it?', 'N/A')}
            Box Includes: {product_info.get('Box Includes', 'N/A')}
            Warranty: {product_info.get('Warranty', 'N/A')}
        """
        else:
            # Beauty products - original FAQ format
            faqs_prompt = f"""
            Write 6 frequently asked questions and concise answers for this Amazon product listing.
            Guidelines:
            - Do NOT include questions about certifications, result timelines, or return/refund policies.
            - Avoid questions like "What is it?" or "How do I use it?" since those are covered elsewhere.
            - Focus on practical use, compatibility, safety (general, not certifications), storage, frequency, product care, and common customer concerns.
            - Keep answers under 150 characters.
            - Use clear, customer-friendly language.
            - Naturally include relevant keywords where possible, but do not keyword stuff.
            - Do not mention price, promotions, or other brands.

            Product: {product_info.get('Product Name', 'N/A')}
            Key Features: {product_info.get('USPs Front', 'N/A')}
            Ingredients: {product_info.get('Ingredients', 'N/A')}
            How to Use: {product_info.get('How to use it?', 'N/A')}

        """
        st.session_state['faqs'] = generate_section("FAQs", faqs_prompt)
    text_box("FAQs", 'faqs', 200)

    st.session_state.setdefault('reviews', '')
    if st.button("Generate Reviews"):
        # Normalize brand name for consistent detection
        brand_name_raw = product_info.get('Brand Name', '').strip()
        brand_name = normalize_brand_name(brand_name_raw) or selected_brand
        
        # Check brand using normalized name
        is_seoulskin = brand_name == 'Seoulskin'
        is_urban_gabru = brand_name == 'Urban Gabru'
        is_urban_yog = brand_name == 'Urban Yog'
        is_makemeebold = brand_name == 'MakeMeeBold'
        
        # Check if product is unisex from 'Ideal For (Gender)' field
        gender_field = product_info.get('Ideal For (Gender)', '').lower()
        is_unisex = 'unisex' in gender_field or ('men' in gender_field and 'women' in gender_field) or ('male' in gender_field and 'female' in gender_field)
        
        # Determine language distribution based on brand
        if is_seoulskin:
            if is_unisex:
                language_instruction = """
                LANGUAGE & STYLE REQUIREMENTS (SEOULSKIN - PREMIUM ENGLISH FOCUS):
                - MAJORITY (12 reviews) should be in premium English with sophisticated tone
                - MINORITY (3 reviews) can include light Hinglish for authenticity
                - English reviews should emphasize K-beauty, innovation, luxury, premium skincare
                - Keep Hinglish reviews minimal and subtle (e.g., "bohot acha", "mujhe pasand aaya")
                - Focus on premium vocabulary: "radiant", "luminous", "transformative", "innovative"
                
                GENDER DISTRIBUTION (UNISEX PRODUCT):
                - 12 reviews from female customers (use female names)
                - 2 reviews from male customers (use male names)
                - 1 review mentioning product as a gifting option
                """
            else:
                language_instruction = """
                LANGUAGE & STYLE REQUIREMENTS (SEOULSKIN - PREMIUM ENGLISH FOCUS):
                - MAJORITY (10-12 reviews) should be in premium English with sophisticated tone
                - MINORITY (3-5 reviews) can include light Hinglish for authenticity
                - English reviews should emphasize K-beauty, innovation, luxury, premium skincare
                - Keep Hinglish reviews minimal and subtle (e.g., "bohot acha", "mujhe pasand aaya")
                - Focus on premium vocabulary: "radiant", "luminous", "transformative", "innovative"
                - 1 review mentioning product as a gifting option
                """
        elif is_urban_yog:
            if is_unisex:
                language_instruction = """
                LANGUAGE & STYLE REQUIREMENTS (URBAN YOG):
                - 10 reviews in Hinglish with some Gen Z words but simple
                - 5 reviews in English (premium tone)
                
                GENDER DISTRIBUTION (UNISEX PRODUCT):
                - 13 reviews from female customers (use female names)
                - 2 reviews from male customers (use male names)
                - 1 review mentioning product as a gifting option
                """
            else:
                language_instruction = """
                LANGUAGE & STYLE REQUIREMENTS (URBAN YOG):
                - 10 reviews in Hinglish with some Gen Z words but simple
                - 5 reviews in English (premium tone)
                - 1 review mentioning product as a gifting option
                """
        elif is_makemeebold:
            if is_unisex:
                language_instruction = """
                LANGUAGE & STYLE REQUIREMENTS (MAKEMEEBOLD):
                - 10 reviews in English but in simple tone
                - 5 reviews in Hinglish with English words
                
                GENDER DISTRIBUTION (UNISEX PRODUCT):
                - 13 reviews from female customers (use female names)
                - 2 reviews from male customers (use male names)
                - 1 review mentioning product as a gifting option
                """
            else:
                language_instruction = """
                LANGUAGE & STYLE REQUIREMENTS (MAKEMEEBOLD):
                - 10 reviews in English but in simple tone
                - 5 reviews in Hinglish with English words
                - 1 review mentioning product as a gifting option
                """
        elif is_urban_gabru:
            # Urban Gabru is men's brand - no unisex consideration
            language_instruction = """
            LANGUAGE & STYLE REQUIREMENTS (URBAN GABRU - MEN'S BRAND):
            - 10 reviews in Hinglish and simple
            - 5 reviews in English
            - All reviews from male customers (use male names)
            - 1 review mentioning product as a gifting option
            """
        else:
            language_instruction = """
            LANGUAGE & STYLE REQUIREMENTS:
            - Use natural Hinglish like real Indian customers speak (mix of Hindi and English words)
            - Use casual, conversational tone with Hindi words like: bohot, acha, laga, pasand, mujhe, thoda, ekdum, zyada, etc.
            - Hindi words should be written in Roman script (Hinglish)
            - Mix both English and Hindi naturally in the same sentence
            - Use common phrases like: "mujhe pasand aaya", "bohot acha h", "thoda zyada", "ekdum fresh"
            - 1 review mentioning product as a gifting option
            """
        
        review_prompt = f"""
            Write 15 customer reviews for this product.
            
            {language_instruction}
            
            REVIEW STRUCTURE:
            - First 2 reviews: Long, story-style testimonials
            - Remaining 13 reviews: Short reviews highlighting different aspects
            - All reviewer names should be common Indian names (first name + last name)
            
            REVIEW CONTENT GUIDELINES:
            - Include answers/benefits that customers generally want to hear
            - Clear doubts and highlight product benefits
            - Use comparisons like "better than [old product/old technology]"
            - Mention relatable problems like "This product really helps because I have [common related problem]"
            - Make reviews authentic and helpful for potential buyers
            
            EXAMPLES OF HINGLISH STYLE TO FOLLOW:
            - "Ye eye patches bohot ache hai, ekdum cooling effect deta h aur eyes fresh lagti h"
            - "First time use kiya aur dark circles thode light lag rahe h, mujhe acha laga ye product"
            - "Daily laptop pe kaam karne se meri eyes bohot tired lagti thi, ye laga ke thoda fresh feel hota h"
            - "Feels like an expensive spa treatment! cooling, de-puffing, and brightening all in one."
            - "The texture, the packaging, the results—everything about these eye patches screams luxury."
            - "pehli baar try kiya aur accha laga, mujhe lagta h regular use se aur best result milega"
            - "thanda thanda lagta h lagane ke baad, mujhe toh ye bohot pasand aaya"
            - "Gifted this to my sister and she absolutely loved it!"
            
            TONE VARIATIONS:
            - Mix casual Hinglish reviews with some premium English reviews
            - Some reviews should be very casual (more Hindi), some semi-formal (more English)
            - Include both positive experiences and realistic observations (like "thoda patience chahiye result ke liye")
            
            IMPORTANT: Provide the output in plain text format only. Do NOT use markdown formatting like **, __, or ###. Just use simple text with reviewer name followed by their review separated by " – ".

            Product: {product_info.get('Product Name', 'N/A')}
            Brand: {brand_name}
            Target Audience: {product_info.get('Target Audience', 'N/A')}
            Ideal For (Gender): {product_info.get('Ideal For (Gender)', 'N/A')}
        """
        st.session_state['reviews'] = generate_section("Reviews", review_prompt)
    text_box("Customer Reviews", 'reviews', 300)

    st.session_state.setdefault('brand_info', '')
    if st.button("Generate Brand & Contact Info"):
        # Normalize brand name for consistency
        brand_name_raw = product_info.get('Brand Name', '').strip()
        normalized_brand_info = normalize_brand_name(brand_name_raw) or selected_brand or product_info.get('Brand Name', 'N/A')
        
        brand_info_prompt = f"""
            Generate a professional brand and contact information section for the website footer. 
            Format it in a clean, easy-to-read structure suitable for a website.
            Include:
            - Brand ownership and marketing information
            - Country of origin
            - Customer support contact details
            - Email address
            
            Make it concise, professional, and customer-friendly.

            Brand: {normalized_brand_info}
            Brand Owned & Marketed By: {product_info.get('Brand Owned & Marketed By', 'N/A')}
            Country of Origin: {product_info.get('Country Of Origin', 'N/A')}
            Email: {product_info.get('Email', 'N/A')}
            Contact: {product_info.get('Contact Us', 'N/A')}
        """
        st.session_state['brand_info'] = generate_section("Brand & Contact Info", brand_info_prompt)
    text_box("Brand & Contact Information", 'brand_info', 150)

    # Export or clear
    if st.button("📥 Download Output as Word"):
        try:
            from docx import Document
            from io import BytesIO

            doc = Document()
            doc.add_heading(f"Product Listing Content – {product_info.get('Product Name', 'Product')}", 0)

            sections = [
                ('Product Title', 'title'),
                ('Amazon Bullet Points', 'bullets'),
                ('Amazon Description', 'description'),
                ('Shopify Description', 'shopify'),
                ('Hero Image Prompts', 'hero'),
                ('A+ Image Prompts', 'a_plus'),
                ('Website Bullet Points', 'web_bullets'),
                ('Website Description', 'web_description'),
                ('USP', 'usp'),
                ("What's in the Box?", 'what_you_get'),
                ('How to Use', 'how_to_use'),
                ('FAQs', 'faqs'),
                ('Customer Reviews', 'reviews'),
                ('Brand & Contact Information', 'brand_info'),
                ('Full Website Content', 'web_full')
            ]

            for title, key in sections:
                content = st.session_state.get(key, '')
                if content:
                    doc.add_heading(title, level=1)
                    doc.add_paragraph(content, style='Normal')

            buffer = BytesIO()
            doc.save(buffer)
            buffer.seek(0)
            st.download_button(
                label="Download Word Document",
                data=buffer,
                file_name="product_listing.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            )
        except ImportError:
            st.error("❌ python-docx library is not installed. Please install it using: pip install python-docx")
        except Exception as e:
            st.error(f"❌ Error creating Word document: {str(e)}")
    if st.button("🧹 Clear All"):
        for key in ['title', 'bullets', 'description', 'shopify', 'hero', 'a_plus', 'web_bullets', 'web_description', 'usp', 'what_you_get', 'how_to_use', 'faqs', 'reviews', 'brand_info', 'web_full']:
            st.session_state[key] = ''
        # Clear SEO keywords data
        st.session_state['seo_keywords_data'] = []
        st.session_state['seo_keywords_text'] = ""
        st.success("✅ All content cleared!")
        st.rerun()

    if st.button("📥 Download Output as JSON"):
        output = {k: st.session_state.get(k, '') for k in ['title', 'bullets', 'description', 'shopify', 'hero', 'a_plus', 'web_bullets', 'web_description', 'usp', 'what_you_get', 'how_to_use', 'faqs', 'reviews', 'brand_info', 'web_full']}
        json_str = json.dumps(output, indent=2)
        st.download_button("Download JSON", data=json_str, file_name="listing_content.json", mime="application/json")

else:
    st.info("Please upload your KLD sheet to get started.")
